from __future__ import annotations

import hashlib
from pathlib import Path

from app.domain.models import Document


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def load_documents_from_dir(kb_dir: Path) -> list[Document]:
    if not kb_dir.exists():
        return []

    documents: list[Document] = []
    for path in sorted(p for p in kb_dir.rglob("*") if p.is_file()):
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        text = load_document_text(path)
        if not text.strip():
            continue
        source = str(path.relative_to(kb_dir)).replace("\\", "/")
        doc_id = hashlib.sha256(f"{source}:{path.stat().st_mtime_ns}".encode()).hexdigest()[:24]
        documents.append(
            Document(
                id=doc_id,
                source=source,
                content=text,
                metadata={"extension": path.suffix.lower(), "size": path.stat().st_size},
            )
        )
    return documents


def load_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8")
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    raise ValueError(f"Unsupported document type: {suffix}")

